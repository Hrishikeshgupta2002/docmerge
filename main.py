import os
import re
from typing import List
import logging
from tempfile import NamedTemporaryFile
import wordninja

from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from docx import Document
from docxcompose.composer import Composer

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI(
    title="DocMerge API",
    description="""
    A production-ready API for merging DOCX documents.
    
    ## Features
    
    * **Merge Documents**: Combine multiple DOCX files into a single document with enterprise-grade merging
    * **Preserves Formatting**: Headers, footers, section breaks, and complex layouts are maintained
    
    ## Technology Stack
    
    * FastAPI for high-performance async API
    * docxcompose.Composer for enterprise-grade document merging
    * Docker-based deployment for consistent environments
    * Railway-ready with optimized configuration
    
    ## Endpoints
    
    * `POST /merge-docx/` - Merge multiple DOCX files
    * `GET /` - Health check endpoint
    * `GET /info` - API information and capabilities
    """,
    version="1.0.0",
    contact={
        "name": "DocMerge API",
        "url": "https://docmerge-production.up.railway.app",
    },
    license_info={
        "name": "MIT",
    },
    servers=[
        {
            "url": "https://docmerge-production.up.railway.app",
            "description": "Production server"
        },
        {
            "url": "http://localhost:8080",
            "description": "Local development server"
        }
    ],
    tags_metadata=[
        {
            "name": "Documents",
            "description": "Document merging operations.",
        },
        {
            "name": "Health",
            "description": "Health check and API information endpoints.",
        },
    ],
)

# Configure CORS for production deployment
# Default allowed origins include Railway domain
default_origins = [
    "https://docmerge-production.up.railway.app",
    "http://docmerge-production.up.railway.app",
]

# Allow custom CORS origins via environment variable (comma-separated)
# If CORS_ORIGINS is set to "*", allow all origins
cors_origins_env = os.getenv("CORS_ORIGINS", "")
if cors_origins_env == "*":
    cors_origins = ["*"]
elif cors_origins_env:
    cors_origins = [origin.strip() for origin in cors_origins_env.split(",")]
else:
    cors_origins = default_origins

app.add_middleware(
    CORSMiddleware,
    allow_origins=cors_origins,
    allow_credentials=True,
    allow_methods=["GET", "POST", "PUT", "DELETE", "OPTIONS"],
    allow_headers=["*"],
    expose_headers=["*"],
)

# Validate file type
def validate_docx_file(file: UploadFile):
    """Validate that the uploaded file is a DOCX file"""
    filename_lower = file.filename.lower() if file.filename else ""

    # Check file extension
    if not filename_lower.endswith('.docx'):
        raise HTTPException(status_code=400, detail="Only DOCX files are allowed for merging")

    # Check content type
    allowed_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if file.content_type and file.content_type != allowed_type:
        logger.warning(f"File {file.filename} has unexpected content type: {file.content_type}")


def validate_document_structure(doc: Document, filename: str) -> bool:
    """
    Validate that a Document object has a valid structure before merging
    
    Args:
        doc: Document object to validate
        filename: Original filename for error messages
        
    Returns:
        True if valid, raises HTTPException if invalid
    """
    try:
        # Check if document has a body element
        if doc.element is None or doc.element.body is None:
            raise HTTPException(
                status_code=400,
                detail=f"Invalid document structure in {filename}: Missing document body"
            )
        
        # Check if body has valid XML structure
        if not hasattr(doc.element.body, 'get') or doc.element.body.tag is None:
            raise HTTPException(
                status_code=400,
                detail=f"Invalid document structure in {filename}: Corrupted XML structure"
            )
        
        return True
    except AttributeError as e:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid document structure in {filename}: {str(e)}"
        )


def is_protected_content(text: str) -> bool:
    """
    Check if text contains protected content that should not be modified
    (URLs, emails, product codes, etc.)
    
    Args:
        text: Text to check
        
    Returns:
        True if text contains protected content
    """
    # Check for URLs
    if re.search(r'https?://|www\.', text, re.IGNORECASE):
        return True
    
    # Check for emails
    if '@' in text and re.search(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', text):
        return True
    
    # Check for all caps short words (likely acronyms/codes)
    if re.match(r'^[A-Z]{1,5}$', text.strip()):
        return True
    
    return False


def smart_word_split(text: str) -> str:
    """
    Intelligently split merged words using regex and dictionary-based segmentation.
    Production-hardened approach that handles CamelCase, acronyms, and long merged words.
    
    Args:
        text: Text string to fix
        
    Returns:
        Fixed text with proper spacing
    """
    if not text:
        return text
    
    # Skip protected content (URLs, emails, codes)
    if is_protected_content(text):
        return text
    
    # Performance optimization: Skip wordninja if no long lowercase sequences exist
    # This prevents running expensive segmentation on normal paragraphs
    if not re.search(r'[a-z]{10,}', text):
        # Still apply basic fixes even if no long sequences
        # Step 1: Fix CamelCase (only if suspicious pattern exists)
        if re.search(r'[a-z]{6,}[A-Z]', text):
            text = re.sub(r'([a-z])([A-Z])', r'\1 \2', text)
            text = re.sub(r'([A-Z]+)([A-Z][a-z])', r'\1 \2', text)
        
        # Step 2: Add space between letters and numbers
        text = re.sub(r'([a-zA-Z])(\d)', r'\1 \2', text)
        text = re.sub(r'(\d)([a-zA-Z])', r'\1 \2', text)
        
        # Step 3: Fix punctuation spacing
        text = re.sub(r',([A-Za-z])', r', \1', text)
        text = re.sub(r'\.([A-Za-z])', r'. \1', text)
        
        # Normalize whitespace
        text = re.sub(r'\s{2,}', ' ', text)
        return text.strip()
    
    # Step 1: Fix CamelCase (only if suspicious pattern exists)
    # Only fix if there's 6+ lowercase followed by uppercase (avoids false positives)
    if re.search(r'[a-z]{6,}[A-Z]', text):
        # Split lowercase-to-uppercase transitions
        text = re.sub(r'([a-z])([A-Z])', r'\1 \2', text)
        # Split acronym-to-word transitions (e.g., "APPDto" -> "APPD to")
        text = re.sub(r'([A-Z]+)([A-Z][a-z])', r'\1 \2', text)
    
    # Step 2: Add space between letters and numbers
    text = re.sub(r'([a-zA-Z])(\d)', r'\1 \2', text)
    text = re.sub(r'(\d)([a-zA-Z])', r'\1 \2', text)
    
    # Step 3: Fix punctuation spacing
    text = re.sub(r',([A-Za-z])', r', \1', text)
    text = re.sub(r'\.([A-Za-z])', r'. \1', text)
    
    # Step 4: Split long lowercase merged words using wordninja
    # Only process words longer than 18 characters to avoid false positives on legitimate compound words
    words = text.split()
    processed_words = []
    
    for word in words:
        # Remove trailing punctuation for processing, then add it back
        trailing_punct = ""
        if word and word[-1] in [',', '.', ';', ':', '!', '?', ')', ']', '}']:
            trailing_punct = word[-1]
            word_clean = word[:-1]
        else:
            word_clean = word
        
        # Only split if word is long (18+ chars), all lowercase, alphabetic, and not a URL
        if (
            len(word_clean) > 18 and
            word_clean.islower() and
            word_clean.isalpha() and
            not word_clean.startswith(("http", "www"))
        ):
            try:
                # Use wordninja to intelligently split the word
                split_words = wordninja.split(word_clean)
                if split_words:
                    processed_words.append(" ".join(split_words) + trailing_punct)
                else:
                    processed_words.append(word)
            except Exception:
                # If wordninja fails, keep original word
                processed_words.append(word)
        else:
            processed_words.append(word)
    
    text = " ".join(processed_words)
    
    # Step 5: Remove double spaces and normalize whitespace
    text = re.sub(r'\s{2,}', ' ', text)
    
    # Step 6: Trim spaces
    text = text.strip()
    
    return text


def normalize_text_spacing(doc: Document):
    """
    Safely normalize text spacing at run level ONLY.
    NEVER reconstructs paragraph text to avoid losing xml:space="preserve" semantics.
    This preserves paragraph XML structure and spacing metadata.
    
    Args:
        doc: Document to process
    """
    try:
        for paragraph in doc.paragraphs:
            # Process each run individually - NEVER reconstruct paragraph text
            # This preserves xml:space="preserve" and spacing metadata
            for run in paragraph.runs:
                if not run.text:
                    continue
                
                original_text = run.text
                
                # Skip protected content
                if is_protected_content(original_text):
                    continue
                
                # Apply smart word splitting to individual run text
                fixed_text = smart_word_split(original_text)
                
                if fixed_text != original_text:
                    run.text = fixed_text
                    logger.debug(f"Fixed spacing in run: '{original_text[:50]}...' -> '{fixed_text[:50]}...'")
    except Exception as e:
        # Don't fail the merge if text normalization fails
        logger.warning(f"Text normalization failed: {str(e)}")
        pass


def ensure_proper_spacing(doc: Document):
    """
    Ensure proper word spacing in paragraphs by adding spaces between text runs where needed.
    This helps prevent words from running together when documents are merged.
    
    Args:
        doc: Document to process
    """
    try:
        for paragraph in doc.paragraphs:
            if len(paragraph.runs) > 1:
                # Process runs to ensure proper spacing
                for i in range(len(paragraph.runs) - 1):
                    current_run = paragraph.runs[i]
                    next_run = paragraph.runs[i + 1]
                    
                    current_text = current_run.text if current_run.text else ""
                    next_text = next_run.text if next_run.text else ""
                    
                    # Skip if either run is empty
                    if not current_text or not next_text:
                        continue
                    
                    # Get last character of current and first character of next
                    last_char = current_text[-1]
                    first_char = next_text[0]
                    
                    # Check if we need to add a space between runs
                    needs_space = False
                    
                    # Case 1: Both are alphanumeric (word characters) - need space
                    if last_char.isalnum() and first_char.isalnum():
                        needs_space = True
                    # Case 2: Last is letter/digit and next starts with letter/digit
                    elif (last_char.isalpha() or last_char.isdigit()) and (first_char.isalpha() or first_char.isdigit()):
                        needs_space = True
                    # Case 3: Last is punctuation that should have space before next word
                    elif last_char in [',', '.', ';', ':', '!', '?'] and first_char.isalnum():
                        needs_space = True
                    
                    # Add space if needed and neither already has whitespace
                    if (needs_space and 
                        not last_char.isspace() and 
                        not first_char.isspace()):
                        # Add space to end of current run
                        current_run.text = current_text + " "
                        logger.debug(f"Added space between text runs: '{current_text[-20:]}' and '{next_text[:20]}'")
    except Exception as e:
        # Don't fail the merge if spacing normalization fails
        logger.warning(f"Could not normalize spacing in document: {str(e)}")
        pass




@app.post("/merge-docx/",
          summary="Merge multiple DOCX files",
          description="Upload multiple DOCX files to merge them into a single document using enterprise-grade merging",
          tags=["Documents"])
async def merge_files(files: List[UploadFile] = File(..., description="List of DOCX files to merge")):
    """
    Merge multiple DOCX files into a single document using docxcompose.Composer
    
    This preserves all formatting including headers, footers, section breaks, and complex layouts.
    
    Note: Documents containing content controls, form fields, or structured document tags may show 
    "[object Object]" placeholders after merging. This is a limitation of the merging library with 
    complex Word document elements. For best results, use documents without content controls.

    Args:
        files: List of DOCX files to merge (minimum 2 files required)

    Returns:
        StreamingResponse: The merged DOCX file
    """
    # Validation: Check if at least 2 files are provided
    if len(files) < 2:
        raise HTTPException(status_code=400, detail="At least 2 files are required for merging")

    # Validation: Check file count limit
    if len(files) > 40:  # Limit to 40 files at once
        raise HTTPException(status_code=400, detail="Maximum 40 files allowed at once")

    # Validate each file - only DOCX allowed
    for file in files:
        validate_docx_file(file)

    docx_files = []
    temp_files = []  # Track all temporary files for cleanup
    output_file = None

    try:
        # Save all DOCX files temporarily for Composer
        for file in files:
            content = await file.read()
            file_extension = file.filename.lower().split('.')[-1]

            if file_extension != 'docx':
                raise HTTPException(status_code=400, detail=f"Only DOCX files are supported. Received: {file_extension}")

            # Save DOCX file temporarily for Composer
            temp_file = NamedTemporaryFile(delete=False, suffix=".docx")
            temp_file.write(content)
            temp_file.close()
            temp_files.append(temp_file.name)
            docx_files.append(temp_file.name)
            logger.debug(f"Added DOCX file: {file.filename}")

        # Merge DOCX files using Composer (preserves formatting, headers, footers, etc.)
        # Validate and load the first DOCX as the master document
        try:
            master_doc = Document(docx_files[0])
            validate_document_structure(master_doc, os.path.basename(docx_files[0]))
            # Normalize spacing in master document before merging
            ensure_proper_spacing(master_doc)
        except Exception as e:
            error_msg = str(e)
            logger.error(f"Error loading master document {docx_files[0]}: {error_msg}")
            if isinstance(e, HTTPException):
                raise
            raise HTTPException(
                status_code=400,
                detail=f"Failed to load master document: {error_msg[:200]}"
            )
        
        composer = Composer(master_doc)

        # Track skipped documents
        skipped_documents = []
        successfully_merged = 1  # Count master document

        # Append remaining DOCX files
        for idx, docx_path in enumerate(docx_files[1:], start=2):
            try:
                # Load and validate document structure before merging
                doc_to_append = Document(docx_path)
                validate_document_structure(doc_to_append, os.path.basename(docx_path))
                
                # Normalize spacing in document before merging to prevent word concatenation
                ensure_proper_spacing(doc_to_append)
                
                # Ensure proper spacing between documents
                # Add a paragraph break to prevent text from running together
                if master_doc.paragraphs:
                    # Check if last paragraph has content - if so, ensure it ends properly
                    last_para = master_doc.paragraphs[-1]
                    if last_para.runs:
                        # Ensure the last run doesn't end without proper spacing
                        last_run = last_para.runs[-1]
                        last_text = last_run.text if last_run.text else ""
                        # If last character is alphanumeric and not whitespace, add a space
                        if last_text and last_text[-1].isalnum() and not last_text[-1].isspace():
                            last_run.text = last_text + " "
                    
                    # Add a blank paragraph to ensure document separation
                    master_doc.add_paragraph()
                
                # Attempt to merge
                composer.append(doc_to_append)
                successfully_merged += 1
                
                logger.debug(f"Merged DOCX: {os.path.basename(docx_path)}")
            except Exception as e:
                error_msg = str(e)
                filename = os.path.basename(docx_path)
                
                # Determine error type for logging
                if "NoneType" in error_msg and "element" in error_msg.lower():
                    error_type = "corrupted structure (NoneType element)"
                elif "multiple relationships" in error_msg and "styles" in error_msg.lower():
                    error_type = "style conflict"
                elif "relationship" in error_msg.lower():
                    error_type = "document structure conflict"
                elif "incorrect type" in error_msg.lower() or "expected" in error_msg.lower():
                    error_type = "invalid document structure"
                else:
                    error_type = "merge error"
                
                # Skip this document and continue
                skipped_documents.append({
                    "document": filename,
                    "position": idx,
                    "error_type": error_type,
                    "error": error_msg[:150]
                })
                logger.warning(f"Skipping document #{idx} ({filename}): {error_type} - {error_msg[:150]}")
                continue

        # Check if we merged at least one document (the master)
        if successfully_merged == 0:
            raise HTTPException(
                status_code=400,
                detail="Failed to merge any documents. All documents appear to be invalid or corrupted."
            )
        
        # Log summary
        if skipped_documents:
            logger.warning(
                f"Merge completed: {successfully_merged} documents merged successfully, "
                f"{len(skipped_documents)} documents skipped"
            )
        else:
            logger.info(f"Merge completed: All {successfully_merged} documents merged successfully")

        # Create output file
        output_file = NamedTemporaryFile(delete=False, suffix=".docx")
        output_file.close()

        # Apply post-merge text normalization ONCE at the end
        # NOTE: This uses run-level fixes ONLY to preserve xml:space="preserve" semantics
        # We NEVER reconstruct paragraph.text to avoid losing spacing metadata
        # Set DISABLE_TEXT_NORMALIZATION=1 to test if this is causing corruption
        disable_normalization = os.getenv("DISABLE_TEXT_NORMALIZATION", "0") == "1"
        
        if not disable_normalization:
            try:
                normalize_text_spacing(master_doc)
                logger.debug("Applied text spacing normalization to fix concatenated words")
            except Exception as e:
                logger.warning(f"Could not normalize spacing: {str(e)}")
                # Continue anyway - spacing normalization is best effort
        else:
            logger.info("Text normalization disabled for testing (DISABLE_TEXT_NORMALIZATION=1)")
        
        # Save the merged document
        master_doc.save(output_file.name)
        logger.info(
            f"Merged document saved to: {output_file.name} "
            f"(merged {successfully_merged}/{len(docx_files)} DOCX files)"
        )

        # Read the merged file content before cleanup
        with open(output_file.name, 'rb') as f:
            content = f.read()

        # Clean up temporary files early to free resources
        for temp_path in temp_files:
            try:
                os.unlink(temp_path)
            except OSError:
                pass  # Ignore errors when deleting temp files

        if output_file and os.path.exists(output_file.name):
            try:
                os.unlink(output_file.name)
            except OSError:
                pass  # Ignore errors when deleting temp files

        # Return the merged file content as a streaming response
        def iterfile():
            yield content

        # Build response headers
        headers = {
            "Content-Disposition": "attachment; filename=merged_document.docx"
        }
        
        # Add warning header if documents were skipped
        if skipped_documents:
            skipped_info = f"{len(skipped_documents)} documents skipped"
            headers["X-Skipped-Documents"] = skipped_info
            logger.info(f"Response includes warning: {skipped_info}")

        return StreamingResponse(
            iterfile(),
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers=headers
        )

    except HTTPException:
        # Re-raise HTTP exceptions (these already have user-friendly messages)
        raise
    except Exception as e:
        error_msg = str(e)
        logger.error(f"Unexpected error during merging: {error_msg}")

        # Ensure cleanup even if there's an exception
        for temp_path in temp_files:
            try:
                if os.path.exists(temp_path):
                    os.unlink(temp_path)
            except OSError:
                pass  # Ignore errors when deleting temp files

        if output_file and os.path.exists(output_file.name):
            try:
                os.unlink(output_file.name)
            except OSError:
                pass  # Ignore errors when deleting temp files

        # Provide more context for common errors
        if "relationship" in error_msg.lower():
            raise HTTPException(
                status_code=500,
                detail=(
                    "Document structure error during merging. "
                    "This may occur when documents have incompatible internal structures or style definitions. "
                    "Please ensure all documents are valid DOCX files."
                )
            )
        else:
            raise HTTPException(
                status_code=500,
                detail=f"An unexpected error occurred during document merging: {error_msg[:300]}"
            )


@app.get("/",
         summary="API Health Check",
         description="Check if the API is running properly",
         tags=["Health"])
async def health_check():
    """Health check endpoint to verify API is running"""
    return {"status": "healthy", "message": "DocMerge API is running"}


@app.get("/info",
         summary="API Information",
         description="Get information about the API capabilities",
         tags=["Health"])
async def api_info():
    """Endpoint to get API information"""
    return {
        "title": "DocMerge API",
        "version": "1.0.0",
        "description": "API for merging DOCX files",
        "features": [
            "Merge multiple DOCX files into a single document",
            "Enterprise-grade DOCX merging using docxcompose.Composer (preserves headers, footers, section breaks)",
            "Handles large documents efficiently",
            "Secure temporary file handling with automatic cleanup",
            "Docker-based deployment with consistent environment"
        ],
        "limitations": [
            "Documents with content controls or form fields may show '[object Object]' placeholders after merging",
            "Complex Word elements (structured document tags) may not be fully preserved",
            "For best results, use standard DOCX files without content controls"
        ],
        "endpoints": {
            "merge": {
                "path": "/merge-docx/",
                "method": "POST",
                "description": "Merge multiple DOCX files",
                "params": "Multiple DOCX files as form data",
                "requirements": "At least 2 DOCX files, max 40 files"
            }
        }
    }